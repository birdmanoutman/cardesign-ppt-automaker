import jieba.posseg as pseg
from collections import Counter
from docx import Document
import re
import jieba

filePath = r'C:\Users\dell\Desktop\share\Syncdisk\TXT\231225_附件三：研发总院2023年度优秀团队奖申请表(2).docx'

def read_text(filePath):
    if filePath.endswith('txt'):
        with open(filePath, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    elif filePath.endswith('docx'):
        doc = Document(filePath)
        full_text = []

        # 提取段落中的文本
        for para in doc.paragraphs:
            full_text.append(para.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)

text = read_text(filePath)

# 向jieba词典中添加新词
custom_words = ['MG', '设计战略', '荣威', '飞凡', '智己', '上汽集团', '研发总院', '设计中心']
for word in custom_words:
    jieba.add_word(word, tag='n')

# 使用jieba进行中文分词和词性标注
words = pseg.cut(text)

# 提取中文名词
cn_nouns = [word for word, flag in words if flag.startswith('n')]

# 使用正则表达式匹配英文单词
en_words = re.findall(r'\b[A-Za-z]+\b', text)

# 假设所有英文单词都算作名词（这可能不完全准确）
en_nouns = en_words

# 合并中文名词和英文名词列表
all_nouns = cn_nouns + en_nouns

# 统计名词出现频率
nouns_freq = Counter(all_nouns)

# 获取出现频率最高的30个名词
top_30_nouns = nouns_freq.most_common(30)

# 打印结果
for noun, freq in top_30_nouns:
    print(f"{noun}: {freq}")

# # 调试信息
# print(f"提取的所有中文名词: {cn_nouns}")
# print(f"提取的所有英文单词: {en_words}")
# print(f"合并后的所有名词: {all_nouns}")
# print(f"名词频率统计: {nouns_freq}")
